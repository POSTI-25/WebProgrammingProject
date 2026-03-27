"""Simple DOCX/DOC text extraction utilities."""

from io import BytesIO

from docx import Document


def extract_doc_text(file_bytes: bytes, extension: str) -> str:
    """
    Extract raw plain text from DOCX bytes.

    For legacy .doc files, this function raises an error so the API can
    return a clear extraction failure response.
    """
    ext = extension.lower()
    if ext == ".doc":
        raise ValueError("Legacy .doc extraction is not supported by python-docx")

    doc = Document(BytesIO(file_bytes))

    paragraphs_text = [para.text for para in doc.paragraphs if para.text.strip()]

    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                tables_text.append(" | ".join(row_text))

    full_text = "\n".join(paragraphs_text)
    if tables_text:
        full_text += "\n\n" + "\n".join(tables_text)

    return full_text